"""
parser.py
---------
Production-grade resume and JD file parser.

Improvements:
  - pdfplumber as primary PDF engine (handles tables, multi-column layouts)
  - PyPDF2 as automatic fallback
  - DOCX: extracts headers, tables, and text boxes (None-safe style check)
  - Encoding detection for TXT files (UTF-8, UTF-16, Latin-1)
  - File size validation (max 10 MB)

Author: SmartHire AI
"""

import io
import logging
from pathlib import Path
from typing import Optional, Union

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_MB    = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


# -- PDF Extraction ---------------------------------------------------

def extract_text_from_pdf(file: Union[bytes, io.BytesIO]) -> str:
    """
    Extract text from PDF using pdfplumber (primary) with PyPDF2 fallback.
    """
    if isinstance(file, bytes):
        file = io.BytesIO(file)

    # Try pdfplumber first
    try:
        import pdfplumber
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            pages_text = []
            for page_num, page in enumerate(pdf.pages):
                page_text  = page.extract_text(x_tolerance=3, y_tolerance=3)
                tables     = page.extract_tables()
                table_text = ""
                for table in tables:
                    for row in table:
                        row_cells  = [cell.strip() if cell else "" for cell in row]
                        table_text += " ".join(row_cells) + "\n"

                combined = ""
                if page_text:
                    combined += page_text
                if table_text:
                    combined += "\n" + table_text
                if combined.strip():
                    pages_text.append(combined)
                else:
                    logger.warning(f"Page {page_num + 1}: no text extracted")

            full_text = "\n\n".join(pages_text)
            if full_text.strip():
                logger.info(f"pdfplumber: extracted {len(full_text)} chars from {len(pdf.pages)} pages")
                return full_text

    except ImportError:
        logger.warning("pdfplumber not installed -- trying PyPDF2")
    except Exception as e:
        logger.warning(f"pdfplumber failed ({e}) -- trying PyPDF2")

    # Fallback: PyPDF2
    try:
        import PyPDF2
        if isinstance(file, io.BytesIO):
            file.seek(0)
        else:
            file = io.BytesIO(file)

        reader     = PyPDF2.PdfReader(file)
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            else:
                logger.warning(f"PyPDF2: no text on page {page_num + 1}")

        full_text = "\n".join(text_parts)
        if full_text.strip():
            logger.info(f"PyPDF2: extracted {len(full_text)} chars")
            return full_text

        raise ValueError("No text extracted from PDF. File may be image-based (scanned).")

    except ImportError:
        raise ImportError("No PDF parser installed. Run: pip install pdfplumber PyPDF2")
    except Exception as e:
        raise ValueError(f"PDF parsing failed: {e}")


# -- DOCX Extraction --------------------------------------------------

def extract_text_from_docx(file: Union[bytes, io.BytesIO]) -> str:
    """
    Extract text from DOCX including paragraphs, tables, and headers.
    Handles None styles safely (some DOCX files have unstyled paragraphs).
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Run: pip install python-docx")

    if isinstance(file, bytes):
        file = io.BytesIO(file)

    try:
        doc   = Document(file)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # FIX: para.style or para.style.name can be None in some DOCX files
            try:
                style_name = para.style.name if para.style and para.style.name else ""
            except Exception:
                style_name = ""

            if style_name.startswith("Heading"):
                parts.append(f"\n{text.upper()}\n")
            else:
                parts.append(text)

        # Extract table contents
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip() if cell.text else ""
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    parts.append(" | ".join(row_cells))

        full_text = "\n".join(parts)

        # Fallback: try reading body XML if no text found
        if not full_text.strip():
            try:
                import re
                xml_content = doc.element.body.xml
                clean = re.sub(r'<[^>]+>', ' ', xml_content)
                clean = re.sub(r'\s+', ' ', clean).strip()
                if clean:
                    logger.warning("DOCX: used XML fallback extraction")
                    return clean
            except Exception:
                pass
            raise ValueError("No text extracted from DOCX — file may be empty or image-based.")

        logger.info(f"DOCX: extracted {len(full_text)} chars")
        return full_text

    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")


# -- TXT Extraction ---------------------------------------------------

def extract_text_from_txt(file: Union[bytes, io.BytesIO, str]) -> str:
    """
    Extract text from TXT file with encoding detection.
    Tries UTF-8, UTF-16, Latin-1 in order.
    """
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

    try:
        if isinstance(file, str):
            for enc in encodings:
                try:
                    with open(file, "r", encoding=enc, errors="strict") as f:
                        text = f.read()
                    logger.info(f"TXT: read {len(text)} chars (encoding: {enc})")
                    return text
                except (UnicodeDecodeError, UnicodeError):
                    continue
            with open(file, "r", encoding="utf-8", errors="replace") as f:
                return f.read()

        elif isinstance(file, bytes):
            raw = file
        elif isinstance(file, io.BytesIO):
            raw = file.read()
        else:
            raise ValueError(f"Unsupported type: {type(file)}")

        for enc in encodings:
            try:
                text = raw.decode(enc)
                logger.info(f"TXT: decoded {len(text)} chars (encoding: {enc})")
                return text
            except (UnicodeDecodeError, UnicodeError):
                continue

        text = raw.decode("utf-8", errors="replace")
        logger.warning("TXT decoded with replacement characters")
        return text

    except Exception as e:
        raise ValueError(f"TXT parsing failed: {e}")


# -- Public API -------------------------------------------------------

def validate_file_size(file_bytes: bytes, filename: str) -> None:
    """Raise ValueError if file exceeds MAX_FILE_SIZE_MB."""
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(
            f"File '{filename}' is {size_mb:.1f} MB — maximum is {MAX_FILE_SIZE_MB} MB."
        )


def parse_resume(file: Union[bytes, io.BytesIO], filename: str) -> str:
    """
    Parse a resume file and return extracted text.
    Supports PDF, DOCX, TXT.
    """
    if isinstance(file, bytes):
        validate_file_size(file, filename)

    suffix = Path(filename).suffix.lower()
    logger.info(f"Parsing resume: {filename} ({suffix})")

    if suffix == ".pdf":
        return extract_text_from_pdf(file)
    elif suffix == ".docx":
        return extract_text_from_docx(file)
    elif suffix == ".txt":
        return extract_text_from_txt(file)
    else:
        raise ValueError(f"Unsupported format: '{suffix}'. Supported: PDF, DOCX, TXT.")


def parse_job_description(
    text_or_file: Union[str, bytes, io.BytesIO],
    filename: Optional[str] = None,
) -> str:
    """
    Parse a job description from pasted text or uploaded file.
    """
    if isinstance(text_or_file, str):
        if not text_or_file.strip():
            raise ValueError("Job description text is empty.")
        logger.info(f"JD received as text ({len(text_or_file)} chars)")
        return text_or_file

    if filename is None:
        raise ValueError("filename is required when passing a file.")

    return parse_resume(text_or_file, filename)
